import os
import asyncio
import logging
from typing import List, Dict, Any, Optional, Union
from pathlib import Path
import hashlib
import json
import tempfile
import re 

import google.generativeai as genai

from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter # Added for more reliable chunking
from langchain_core.documents.base import Document
import tiktoken
from supabase import create_client, Client
import PyPDF2
from docx import Document as DocxDocument
from ..core.gemini_key_manager import get_key_manager, safe_embed_content



# Import new config file
from ..config.rag_config import (
    get_embedding_config,
    get_chunk_config,
    get_database_config,
    get_performance_config
)

# Import vector utilities
from ..utils.vector_utils import ensure_768_dimensions, log_vector_info


key_manager = get_key_manager()

logger = logging.getLogger(__name__)

class DocumentProcessor:
    def __init__(self):
        logger.debug("Initializing DocumentProcessor...")
        
        # Get environment variables with proper error handling
        supabase_url = os.getenv("SUPABASE_URL")
        supabase_key = os.getenv("SUPABASE_SERVICE_KEY") or os.getenv("SUPABASE_KEY")
        
        if not supabase_url:
            raise ValueError("SUPABASE_URL environment variable is required")
        if not supabase_key:
            raise ValueError("SUPABASE_SERVICE_KEY or SUPABASE_KEY environment variable is required")
            
        self.supabase: Client = create_client(supabase_url, supabase_key)
        logger.debug(f"Supabase client initialized for URL: {supabase_url}")
        
        # Get settings from new config
        self.embedding_config = get_embedding_config()
        self.chunk_config = get_chunk_config()
        self.db_config = get_database_config()
        self.performance_config = get_performance_config()
        
        api_key = os.getenv("GEMINI_API_KEY")
        if not api_key:
            logger.warning("GEMINI_API_KEY not found in environment variables during DocumentProcessor init.")
            raise ValueError("GEMINI_API_KEY environment variable is required")
        else:
            logger.info("GEMINI_API_KEY found in environment during DocumentProcessor init.")
            
        # Configure genai properly - FIXED
        try:
            genai.configure(api_key=api_key)
            logger.debug("genai configured with API key.")
        except AttributeError:
            # Alternative configuration method if configure not available
            os.environ["GOOGLE_API_KEY"] = api_key
            logger.debug("GOOGLE_API_KEY set in environment as fallback.")
        
        logger.debug("Initializing text splitters...")
        
        # Primary text splitter: RecursiveCharacterTextSplitter
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.chunk_config.DEFAULT_CHUNK_SIZE,
            chunk_overlap=self.chunk_config.DEFAULT_CHUNK_OVERLAP,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""]
        )
        
        logger.debug(f"Text splitter initialized with chunk_size={self.chunk_config.DEFAULT_CHUNK_SIZE}, "
                    f"overlap={self.chunk_config.DEFAULT_CHUNK_OVERLAP}")
        
        self.encoding = tiktoken.get_encoding("cl100k_base")
        logger.debug("tiktoken encoding cl100k_base loaded.")
        
        self.enhanced_processor = None
        
        logger.info(f"DocumentProcessor initialized successfully with config - "
                   f"Max chunks per doc: {self.chunk_config.MAX_CHUNKS_PER_DOCUMENT}, "
                   f"Chunk size: {self.chunk_config.DEFAULT_CHUNK_SIZE}")

    async def process_document(self, document_id: int, file_path: str) -> Dict[str, Any]:
        """
        Processes a document: splits into chunks, generates embeddings with contextual headers,
        and saves them to the database including metadata for references.
        """
        logger.info(f"Starting process_document for ID: {document_id}, Path: {file_path}")
        try:
            logger.debug(f"Fetching document name for ID: {document_id}")
            doc_result = self.supabase.table(self.db_config.DOCUMENTS_TABLE).select("name").eq("id", document_id).execute()
            if doc_result.data and doc_result.data[0].get("name"):
                document_name = doc_result.data[0]["name"]
                logger.debug(f"Document name '{document_name}' found for ID: {document_id}.")
            else:
                document_name = f"Document {document_id}"
                logger.warning(f"Document name not found for ID: {document_id}, using default: '{document_name}'. Response: {doc_result}")
            
            # Use document_name as the initial document title / header
            current_overall_header = document_name 
            
            logger.debug(f"Updating document status to 'processing' for ID: {document_id}")
            await self._update_document_status(document_id, "processing")
            
            # Load and split document using the configured text_splitter
            logger.debug(f"Calling _load_and_split_document for file: {file_path}")
            raw_chunks = await self._load_and_split_document(file_path)
            logger.info(f"Document ID: {document_id} ('{document_name}') split into {len(raw_chunks)} raw chunks using {type(self.text_splitter).__name__}.")
            
            # Check if number of chunks exceeds limit from config
            if len(raw_chunks) > self.chunk_config.MAX_CHUNKS_PER_DOCUMENT:
                logger.warning(f"Number of chunks ({len(raw_chunks)}) exceeds maximum limit "
                             f"({self.chunk_config.MAX_CHUNKS_PER_DOCUMENT}) for document ID: {document_id}")
                raw_chunks = raw_chunks[:self.chunk_config.MAX_CHUNKS_PER_DOCUMENT]
                logger.info(f"Trimmed chunks to {len(raw_chunks)} for document ID: {document_id}")
            
            # Ensure API key is configured for embedding generation
            api_key = os.environ.get("GEMINI_API_KEY")
            if not api_key:
                logger.error(f"GEMINI_API_KEY not found for document ID: {document_id} processing.")
                await self._update_document_status(document_id, "failed", note="Missing GEMINI_API_KEY")
                return {"success": False, "error": "Missing GEMINI_API_KEY", "document_id": document_id}

            processed_chunks_for_db = []
            chunk_meta_info_for_bc = []

            for i, chunk_doc in enumerate(raw_chunks):
                try:
                    if not chunk_doc.page_content.strip():
                        logger.warning(f"Empty content in chunk {i} from Langchain Splitter (doc ID: {document_id}), skipping.")
                        continue
                        
                    # Determine header and page number from metadata if available
                    header_to_use = chunk_doc.metadata.get('header', f"Chunk {i+1}")
                    page_number = chunk_doc.metadata.get('page_number')

                    # Original text (without header) for token count
                    original_chunk_text = chunk_doc.page_content
                    
                    # Text to be embedded and stored (potentially with header prepended)
                    chunk_text_with_header = f"{header_to_use}\n\n{original_chunk_text}" if header_to_use else original_chunk_text

                    # Check if we've reached the vector limit during processing
                    if len(processed_chunks_for_db) >= self.chunk_config.MAX_CHUNKS_PER_DOCUMENT:
                        logger.warning(f"Maximum number of vectors ({self.chunk_config.MAX_CHUNKS_PER_DOCUMENT}) "
                                     f"reached during processing for document ID: {document_id}")
                        break

                    embedding_vector = await self._generate_embedding(chunk_text_with_header)
                    if not embedding_vector or len(embedding_vector) == 0:
                        logger.error(f"Empty embedding returned for chunk {i} (doc ID: {document_id}). "
                                   f"Original text snippet: {original_chunk_text[:100]}...")
                        continue
                    
                    # Token count using configured encoding
                    token_count = len(self.encoding.encode(original_chunk_text))
                    
                    # Check token limits from config
                    if token_count > self.chunk_config.MAX_TOKENS_PER_CHUNK:
                        logger.warning(f"Chunk {i} exceeds max token limit: {token_count} > {self.chunk_config.MAX_TOKENS_PER_CHUNK}")
                    
                    if token_count < self.chunk_config.MIN_TOKENS_PER_CHUNK:
                        logger.warning(f"Chunk {i} below min token limit: {token_count} < {self.chunk_config.MIN_TOKENS_PER_CHUNK}")

                    processed_chunks_for_db.append({
                        "document_id": document_id,
                        "chunk_text": chunk_text_with_header,
                        "embedding": embedding_vector,
                        "content_token_count": token_count
                    })
                    
                    chunk_meta_info_for_bc.append({
                        "chunk_index": i, 
                        "original_text": original_chunk_text, 
                        "header": header_to_use
                    })
                    
                    logger.info(f"Generated embedding for chunk {i+1}/{len(raw_chunks)} "
                              f"(doc ID: {document_id}, tokens: {token_count})")
                              
                except Exception as e_emb:
                    logger.error(f"Error generating embedding or preparing chunk {i} (doc ID: {document_id}): {e_emb}", exc_info=True)
                    continue
            
            if not processed_chunks_for_db:
                logger.warning(f"No processable chunks with embeddings found for document {document_id}")
                await self._update_document_status(document_id, "failed", "No processable chunks with embeddings after generation.")
                return {"success": False, "document_id": document_id, "chunks_created": 0, 
                       "error": "No processable chunks with embeddings after generation"}

            # Save chunks to database using RPC
            logger.info(f"Attempting to save {len(processed_chunks_for_db)} processed chunks to "
                       f"'{self.db_config.CHUNKS_TABLE}' via RPC for doc ID: {document_id}.")
            
            successful_rpc_inserts = 0
            any_rpc_insert_failed = False
            first_rpc_error_message = ""

            for i, chunk_data_to_insert in enumerate(processed_chunks_for_db):
                try:
                    logger.debug(f"Preparing to insert chunk {i+1}/{len(processed_chunks_for_db)} via RPC for doc ID: {document_id}. "
                                f"Snippet: '{chunk_data_to_insert['chunk_text'][:50]}...', "
                                f"Emb Dim: {len(chunk_data_to_insert['embedding']) if chunk_data_to_insert.get('embedding') else 'N/A'}")

                    params_for_rpc = {
                        "p_document_id": chunk_data_to_insert["document_id"],
                        "p_chunk_text": chunk_data_to_insert["chunk_text"],
                        "p_embedding": chunk_data_to_insert["embedding"],
                        "p_content_token_count": chunk_data_to_insert["content_token_count"]
                    }
                    
                    rpc_response = self.supabase.rpc("insert_document_chunk_basic", params_for_rpc).execute()

                    if hasattr(rpc_response, 'data') and rpc_response.data and len(rpc_response.data) > 0:
                        inserted_chunk_id = rpc_response.data[0].get('id', 'N/A') # RPC returns the inserted row
                        logger.info(f"Successfully inserted chunk {i+1}/{len(processed_chunks_for_db)} via RPC for doc ID: {document_id}. Inserted '{self.db_config.CHUNKS_TABLE}' ID: {inserted_chunk_id}")
                        successful_rpc_inserts += 1
                    else:
                        # Handle error cases more robustly
                        error_msg = f"RPC call for '{self.db_config.CHUNKS_TABLE}' chunk {i+1} (doc ID: {document_id}) returned no data"
                        logger.error(error_msg)
                        any_rpc_insert_failed = True
                        if not first_rpc_error_message: 
                            first_rpc_error_message = error_msg
                except Exception as e_rpc_call:
                    error_msg = f"Exception during RPC call for '{self.db_config.CHUNKS_TABLE}' chunk {i+1} (doc ID: {document_id}): {e_rpc_call}"
                    logger.error(error_msg, exc_info=True)
                    any_rpc_insert_failed = True
                    if not first_rpc_error_message: first_rpc_error_message = error_msg
            
            logger.info(f"Finished RPC inserts for '{self.db_config.CHUNKS_TABLE}' for doc ID: {document_id}. Successfully inserted: {successful_rpc_inserts}/{len(processed_chunks_for_db)}.")

            if any_rpc_insert_failed or (len(processed_chunks_for_db) > 0 and successful_rpc_inserts == 0):
                final_status_reason = first_rpc_error_message or "One or more chunks failed to insert into '{self.db_config.CHUNKS_TABLE}' via RPC."
                logger.error(f"Overall '{self.db_config.CHUNKS_TABLE}' RPC insertion for doc ID {document_id} failed or partially failed. Error: {final_status_reason}")
                await self._update_document_status(document_id, "failed", final_status_reason)
                return {"success": False, "document_id": document_id, "chunks_created": successful_rpc_inserts, "error": final_status_reason}
            
            # If we reach here, all RPC inserts to '{self.db_config.CHUNKS_TABLE}' were successful.
            logger.info(f"All {successful_rpc_inserts} chunks for document ID {document_id} saved successfully to '{self.db_config.CHUNKS_TABLE}' via RPC.")
            
            # --- Backward Compatibility: Save to 'embeddings' table ---
            # This uses 'chunk_meta_info_for_bc' and 'processed_chunks_for_db' which should be aligned.
            if successful_rpc_inserts > 0: # Only proceed if primary chunks were saved
                logger.info(f"Attempting to save {len(chunk_meta_info_for_bc)} items to 'embeddings' table for backward compatibility (doc ID: {document_id}).")
                bc_inserts_count = 0
                for i, meta_info in enumerate(chunk_meta_info_for_bc):
                    # Find the corresponding processed chunk data for the embedding
                    # This assumes processed_chunks_for_db[i] corresponds to chunk_meta_info_for_bc[i]
                    if i < len(processed_chunks_for_db):
                        embedding_vector_for_bc = processed_chunks_for_db[i]["embedding"]
                        
                        embedding_data_for_bc_table = {
                            "content": meta_info["original_text"], 
                            "metadata": {"document_id": document_id, "chunk_index": meta_info["chunk_index"], "header": meta_info.get("header")},
                            "embedding": embedding_vector_for_bc,
                        }
                        try:
                            bc_response = self.supabase.table("embeddings").insert(embedding_data_for_bc_table).execute()
                            if hasattr(bc_response, 'data') and bc_response.data:
                                bc_inserts_count += 1
                            else:
                                logger.error(f"BC Save: Error inserting item {i+1} to 'embeddings' for doc {document_id}")
                        except Exception as e_bc_insert_exc:
                            logger.error(f"BC Save: Exception inserting item {i+1} to 'embeddings' for doc {document_id}: {e_bc_insert_exc}", exc_info=True)
                    else:
                        logger.warning(f"BC Save: Mismatch between chunk_meta_info_for_bc and processed_chunks_for_db at index {i}. Cannot save to 'embeddings' table.")
                
                logger.info(f"Backward compatibility: Saved {bc_inserts_count}/{len(chunk_meta_info_for_bc)} items to 'embeddings' table for doc ID: {document_id}.")
            
            # Update main document status and content (if all primary chunks saved)
            await self._update_document_status(document_id, "completed")
            
            # Only try to update document content if we processed fewer than MAX_VECTORS_PER_DOCUMENT
            if len(raw_chunks) <= self.chunk_config.MAX_CHUNKS_PER_DOCUMENT:
                await self._update_document_content(document_id, raw_chunks)
            else:
                logger.warning(f"Skipping document content update due to large chunk count for document ID: {document_id}")
            
            return {
                "success": True,
                "document_id": document_id,
                "chunks_created": successful_rpc_inserts,
                "total_chunks": len(raw_chunks),
                "trimmed": len(raw_chunks) > self.chunk_config.MAX_CHUNKS_PER_DOCUMENT
            }
            
        except Exception as e:
            logger.error(f"Error in process_document for ID {document_id}: {str(e)}", exc_info=True)
            await self._update_document_status(document_id, "failed", str(e))
            return {"success": False, "document_id": document_id, "chunks_created": 0, "error": str(e)}

    async def _load_and_split_document(self, file_path: str) -> List[Document]:
        """
        Loads a document from file_path and splits it into chunks.
        Returns list of Document objects with metadata.
        """
        logger.debug(f"Starting _load_and_split_document for file: {file_path}")
        
        try:
            # Basic metadata to include with each chunk
            original_metadata = {
                "source": file_path,
                "file_name": Path(file_path).name
            }
            
            # Extract text based on file type
            file_extension = Path(file_path).suffix.lower()
            if file_extension == '.pdf':
                doc_text_content = self._extract_pdf_text(file_path)
            elif file_extension in ['.docx', '.doc']:
                doc_text_content = self._extract_docx_text(file_path)
            elif file_extension in ['.txt', '.md', '.html', '.json']:
                with open(file_path, 'r', encoding='utf-8') as f:
                    doc_text_content = f.read()
                logger.debug(f"Read {len(doc_text_content)} characters from text file: {file_path}")
            else:
                logger.error(f"Unsupported file type: {file_extension} for file {file_path}")
                raise ValueError(f"Unsupported file type: {file_extension}")
            
            if not doc_text_content.strip():
                logger.warning(f"Document {file_path} has no text content after extraction. Returning empty list of chunks.")
                return []

            logger.debug(f"Splitting text content (length: {len(doc_text_content)}) using {type(self.text_splitter).__name__}'s split_documents method...")
            
            # Create a proper Document object using langchain's Document class
            doc = Document(page_content=doc_text_content, metadata=original_metadata)
            
            # Use split_documents which returns List[Document]
            chunks = self.text_splitter.split_documents([doc])
            logger.info(f"Successfully split document {file_path} into {len(chunks)} chunks.")

            # Add metadata for each chunk
            for i, chunk in enumerate(chunks):
                chunk.metadata["chunk_index"] = i
                
            return chunks
            
        except Exception as e:
            logger.error(f"Error in _load_and_split_document for {file_path}: {str(e)}", exc_info=True)
            raise

    def _extract_pdf_text(self, file_path: str) -> str:
        """Extracts text from PDF file"""
        logger.debug(f"Starting _extract_pdf_text for: {file_path}")
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                num_pages = len(pdf_reader.pages)
                logger.debug(f"PDF has {num_pages} pages.")
                for i, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                    logger.debug(f"Extracted text from page {i+1}/{num_pages}. Length: {len(page_text) if page_text else 0}")
            logger.info(f"Finished extracting text from PDF: {file_path}. Total length: {len(text)}")
        except Exception as e:
            logger.error(f"Error extracting PDF text from {file_path}: {str(e)}", exc_info=True)
            raise
        return text

    def _extract_docx_text(self, file_path: str) -> str:
        """Extracts text from DOCX file"""
        logger.debug(f"Starting _extract_docx_text for: {file_path}")
        try:
            doc = DocxDocument(file_path)
            text = ""
            logger.debug(f"DOCX file has {len(doc.paragraphs)} paragraphs.")
            for i, paragraph in enumerate(doc.paragraphs):
                text += paragraph.text + "\n"
                # Add less verbose logging here, maybe every N paragraphs or by text length
            logger.info(f"Finished extracting text from DOCX: {file_path}. Total length: {len(text)}")
            return text
        except Exception as e:
            logger.error(f"Error extracting DOCX text from {file_path}: {str(e)}", exc_info=True)
            raise

    async def _generate_embedding(self, text: str, is_query: bool = False) -> Optional[List[float]]:
        """Creates embedding for text"""
        task_type = "retrieval_query" if is_query else "retrieval_document"
        cleaned_text = text[:50].replace('\n', ' ')
        logger.debug(f"Attempting to generate embedding for text (first 50 chars): '{cleaned_text}' with task_type: {task_type}")

        if not text or not text.strip():
            logger.warning(f"Cannot generate embedding for empty or whitespace-only text. Task type: {task_type}")
            return None
        
        # Ensure API key is available (it should be, from __init__ or process_document)
        api_key = os.environ.get("GEMINI_API_KEY")
        if not api_key:
            logger.error("GEMINI_API_KEY not found when trying to generate embedding.")
            return None
        
        try:
            # The model name for embeddings might be different from generative models.
            # Using the model specified for embeddings, e.g., "models/embedding-001"
            # Task types: "retrieval_query", "retrieval_document", "semantic_similarity", "classification", "clustering"
            logger.debug(f"Calling genai.embed_content with model {self.embedding_config.MODEL_NAME} and task_type {task_type}")
            result = await safe_embed_content(
                model=self.embedding_config.MODEL_NAME,
                content=text,
                task_type=task_type
            )
            
            raw_embedding = result["embedding"] if result and 'embedding' in result else None
            if not raw_embedding:
                logger.error(f"No embedding returned from API for task_type: {task_type}")
                return None
            
            # Verify that the vector is exactly 768 dimensions
            embedding = ensure_768_dimensions(raw_embedding)
            
            # Log debug info if there's a problem with the vector size
            if len(raw_embedding) != 768:
                logger.warning(f"Vector dimension adjusted from {len(raw_embedding)} to 768 for task_type: {task_type}")
                log_vector_info(raw_embedding, f"Original {task_type} embedding")
                log_vector_info(embedding, f"Adjusted {task_type} embedding")
            
            logger.debug(f"Successfully generated embedding for task_type: {task_type}. Final embedding length: {len(embedding)}")
            return embedding
            
        except Exception as e:
            logger.error(f"Error generating Gemini embedding for task_type {task_type}: {e}", exc_info=True)
            return None

    async def _update_document_status(self, document_id: int, status: str, note: Optional[str] = None):
        """Updates document status in database"""
        logger.info(f"Updating status for document ID {document_id} to '{status}'. Note: '{note or ''}'")
        try:
            update_data = {"processing_status": status}
            if note is not None:
                update_data["processing_notes"] = note
            
            response = self.supabase.table(self.db_config.DOCUMENTS_TABLE).update(update_data).eq("id", document_id).execute()
            # Check response for errors (Supabase client specific)
            if hasattr(response, 'data') and response.data:
                logger.debug(f"Successfully updated status for document ID {document_id} to {status}. Response: {response.data}")
            else:
                logger.warning(f"Unknown Supabase response or no data returned when updating status for document ID {document_id}: {response}")
        except Exception as e:
            logger.error(f"Exception updating document status for ID {document_id} to {status}: {e}", exc_info=True)

    async def _update_document_content(self, document_id: int, raw_chunks: List[Document]) -> None:
        """Updates document content in database"""
        logger.debug(f"Updating document content for ID: {document_id} with {len(raw_chunks)} chunks")
        try:
            # Create a summary of the document content
            content_summary = "\n".join([chunk.page_content[:200] + "..." if len(chunk.page_content) > 200 else chunk.page_content for chunk in raw_chunks[:5]])
            
            update_data = {
                "content_summary": content_summary,
                "total_chunks": len(raw_chunks)
            }
            
            response = self.supabase.table(self.db_config.DOCUMENTS_TABLE).update(update_data).eq("id", document_id).execute()
            
            if hasattr(response, 'data') and response.data:
                logger.debug(f"Successfully updated content for document ID {document_id}")
            else:
                logger.warning(f"Failed to update content for document ID {document_id}")
                
        except Exception as e:
            logger.error(f"Exception updating document content for ID {document_id}: {e}", exc_info=True)

    async def delete_document_and_all_chunks(self, document_id: int) -> Dict[str, Any]:
        """
        Deletes a document and all its associated chunks from both old and new tables.
        Deletes from: advanced_document_chunks, document_chunks, and documents tables.
        """
        logger.info(f"Attempting to delete document and all its chunks for document_id: {document_id}")
        
        deleted_chunks_count = 0
        deleted_advanced_chunks_count = 0
        try:
            # Delete chunks from 'advanced_document_chunks' table (new system)
            logger.debug(f"Deleting chunks for document_id: {document_id} from 'advanced_document_chunks' table.")
            delete_advanced_chunks_response = self.supabase.table("advanced_document_chunks").delete().eq("document_id", document_id).execute()
            
            if delete_advanced_chunks_response.data:
                deleted_advanced_chunks_count = len(delete_advanced_chunks_response.data)
                logger.info(f"Successfully deleted {deleted_advanced_chunks_count} advanced chunks for document_id: {document_id}.")
            

            logger.debug(f"Deleting chunks for document_id: {document_id} from 'document_chunks' table.")
            delete_chunks_response = self.supabase.table("document_chunks").delete().eq("document_id", document_id).execute()
            
            if delete_chunks_response.data:
                deleted_chunks_count = len(delete_chunks_response.data)
                logger.info(f"Successfully deleted {deleted_chunks_count} old chunks for document_id: {document_id}.")

            # Delete the document from 'documents' table
            logger.debug(f"Deleting document record for document_id: {document_id} from 'documents' table.")
            delete_document_response = self.supabase.table(self.db_config.DOCUMENTS_TABLE).delete().eq("id", document_id).execute()

            document_deleted_successfully = bool(delete_document_response.data)
            
            if document_deleted_successfully:
                logger.info(f"Successfully deleted document record for document_id: {document_id}.")

            return {
                "success": True, 
                "deleted_chunks_count": deleted_chunks_count,
                "deleted_advanced_chunks_count": deleted_advanced_chunks_count,
                "document_deleted": document_deleted_successfully
            }
            
        except Exception as e:
            logger.error(f"Exception in delete_document_and_all_chunks for document_id {document_id}: {e}", exc_info=True)
            return {"success": False, "error": str(e), "deleted_chunks_count": deleted_chunks_count, "document_deleted": False}

    async def delete_document_embeddings(self, document_id: int) -> Dict[str, Any]:
        """
        Deletes all embeddings (chunks) for a given document ID from both old and new systems.
        This is used, for example, before reprocessing a document.
        """
        logger.info(f"Attempting to delete all embeddings/chunks for document_id: {document_id}")
        try:
            # Delete from new advanced_document_chunks table
            logger.info(f"Deleting chunks from 'advanced_document_chunks' for document_id: {document_id}")
            delete_advanced_response = self.supabase.table("advanced_document_chunks").delete().eq("document_id", document_id).execute()
            
            advanced_deleted_count = len(delete_advanced_response.data) if delete_advanced_response.data else 0
            logger.info(f"Deleted {advanced_deleted_count} chunks from 'advanced_document_chunks' for document_id: {document_id}")


            logger.info(f"Deleting chunks from 'document_chunks' for document_id: {document_id}")
            delete_old_response = self.supabase.table("document_chunks").delete().eq("document_id", document_id).execute()
            
            old_deleted_count = len(delete_old_response.data) if delete_old_response.data else 0
            logger.info(f"Deleted {old_deleted_count} chunks from 'document_chunks' for document_id: {document_id}")


            try:
                logger.info(f"Attempting to delete old embeddings from 'embeddings' table for document_id: {document_id}.")
                delete_old_embeddings_response = self.supabase.table("embeddings").delete().eq("document_id", document_id).execute()
                
                embeddings_deleted_count = len(delete_old_embeddings_response.data) if delete_old_embeddings_response.data else 0
                logger.info(f"Deleted {embeddings_deleted_count} entries from 'embeddings' table for document_id: {document_id}")
            except Exception as e_old_emb_del:
                logger.warning(f"Exception during deletion from old 'embeddings' table for document_id {document_id}: {e_old_emb_del}")
            
            return {
                "success": True, 
                "message": "All embeddings/chunks deleted successfully.",
                "advanced_chunks_deleted": advanced_deleted_count,
                "old_chunks_deleted": old_deleted_count
            }
        except Exception as e:
            logger.error(f"General error in delete_document_embeddings for document_id {document_id}: {str(e)}", exc_info=True)
            return {"success": False, "message": str(e)}

    async def search_documents(self, query: str, limit: int = 10, threshold: Optional[float] = None) -> List[Dict[str, Any]]:
        """Semantic search using the new system"""
        logger.info(f"Starting enhanced search with query (first 50 chars): '{query[:50]}', limit: {limit}, threshold: {threshold}")
        try:
            # השתמש בdefault threshold מ-config אם לא סופק
            if threshold is None:
                threshold = self.embedding_config.DEFAULT_SIMILARITY_THRESHOLD
            
            # Enhanced processor is not implemented yet, use fallback
            logger.info("Using fallback search")
            return await self._fallback_search(query, limit, threshold)
            
        except Exception as e:
            logger.error(f"Error during enhanced search: {e}", exc_info=True)
            # Fallback to direct embedding search with default threshold
            fallback_threshold = threshold or self.embedding_config.DEFAULT_SIMILARITY_THRESHOLD
            return await self._fallback_search(query, limit, fallback_threshold)

    async def _fallback_search(self, query: str, limit: int, threshold: float) -> List[Dict[str, Any]]:
        """Fallback search method using direct embedding search"""
        try:
            logger.info("Performing fallback search using direct embedding")
            query_embedding = await self._generate_embedding(query, is_query=True)
            if not query_embedding:
                logger.error("Failed to generate query embedding for fallback")
                return []
            
            response = self.supabase.rpc(
                "advanced_semantic_search",
                {
                    "query_embedding": query_embedding,
                    "similarity_threshold": threshold,
                    "match_count": limit,
                },
            ).execute()
            
            if response.data:
                logger.info(f"Fallback search found {len(response.data)} results")
                return response.data
            return []
                
        except Exception as fallback_e:
            logger.error(f"Fallback search also failed: {fallback_e}", exc_info=True)
            return []

    async def hybrid_search(self, query: str, limit: int = 10, threshold: Optional[float] = None) -> List[Dict[str, Any]]:
        """Hybrid search using the new system"""
        logger.info(f"Starting enhanced hybrid search with query (first 50 chars): '{query[:50]}', limit: {limit}, threshold: {threshold}")
        try:
            # Use default threshold from config if not provided
            if threshold is None:
                threshold = self.embedding_config.DEFAULT_HYBRID_THRESHOLD
            
            # Enhanced processor not implemented yet, use fallback
            logger.info("Enhanced processor not available for hybrid search, using fallback")
            return await self.search_documents(query, limit, threshold)
                
        except Exception as e:
            logger.error(f"Error during enhanced hybrid search: {e}", exc_info=True)
            # Fallback to regular search
            return await self.search_documents(query, limit, threshold)

    async def _get_all_document_chunks(self, limit: int = 100) -> List[Dict[str, Any]]:
        """Get a set of document chunks for local processing"""
        try:
            response = self.supabase.table("document_chunks").select(
                "id,document_id,chunk_text,chunk_header,page_number,section,documents(name)"
            ).limit(limit).execute()
            
            if hasattr(response, 'data') and response.data:
                # Transform the results to match our expected format
                results = []
                for item in response.data:
                    document_name = None
                    if 'documents' in item and item['documents']:
                        document_name = item['documents'].get('name')
                    
                    results.append({
                        'id': item.get('id'),
                        'document_id': item.get('document_id'),
                        'document_name': document_name,
                        'chunk_text': item.get('chunk_text', ''),
                        'chunk_header': item.get('chunk_header', ''),
                        'page_number': item.get('page_number'),
                        'section': item.get('section', '')
                    })
                return results
            return []
        except Exception as e:
            logger.error(f"Error in _get_all_document_chunks: {e}", exc_info=True)
            return [] 