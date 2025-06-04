import os
import asyncio
import logging
from typing import List, Dict, Any, Optional
from pathlib import Path
import hashlib
import json
import tempfile
import re # Import re for regex matching

import google.generativeai as genai
# from langchain.text_splitter import SemanticChunker # Old import
from langchain_experimental.text_splitter import SemanticChunker # New import
from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter # Added for more reliable chunking
import tiktoken
from supabase import create_client, Client
import PyPDF2
from docx import Document as DocxDocument

# Import SemanticChunker and Gemini Embeddings
from langchain_google_genai import GoogleGenerativeAIEmbeddings

logger = logging.getLogger(__name__)

# Constants for chunking limits
MAX_VECTORS_PER_DOCUMENT = 500  # Maximum number of vectors to create per document
DEFAULT_CHUNK_SIZE = 2000        # Size in characters for RecursiveCharacterTextSplitter
DEFAULT_CHUNK_OVERLAP = 200     # Overlap in characters for RecursiveCharacterTextSplitter

class DocumentProcessor:
    def __init__(self):
        logger.debug("Initializing DocumentProcessor...")
        self.supabase: Client = create_client(
            os.getenv("SUPABASE_URL"),
            os.getenv("SUPABASE_SERVICE_KEY") or os.getenv("SUPABASE_KEY")
        )
        logger.debug(f"Supabase client initialized for URL: {os.getenv('SUPABASE_URL')}")
        api_key = os.getenv("GEMINI_API_KEY")
        if not api_key:
            logger.warning("GEMINI_API_KEY not found in environment variables during DocumentProcessor init.")
            # SemanticChunker will fail without embeddings, which likely need the API key.
            # Consider raising an error or having a fallback splitter if API key is missing.
        else:
            logger.info("GEMINI_API_KEY found in environment during DocumentProcessor init.")
            
        genai.configure(api_key=api_key) # This configures the global genai object
        logger.debug("genai configured with API key.")
        
        # Initialize Gemini Embeddings for SemanticChunker
        # Ensure the model name is correct for embeddings.
        # "gemini-pro" is a generative model, not an embedding model.
        # "models/embedding-001" is used later in _generate_embedding.
        logger.debug("Initializing GoogleGenerativeAIEmbeddings...")
        try:
            gemini_embeddings = GoogleGenerativeAIEmbeddings(
                model="models/embedding-001", 
                task_type="retrieval_document",
                google_api_key=api_key  # Explicitly pass the API key
            )
            logger.debug(f"GoogleGenerativeAIEmbeddings initialized successfully with model models/embedding-001.")
        except Exception as e_emb_init:
            logger.error(f"Failed to initialize GoogleGenerativeAIEmbeddings: {e_emb_init}", exc_info=True)
            # Handle or re-raise, as this is critical for SemanticChunker
            raise

        # Initialize both text splitters - we'll use RecursiveCharacterTextSplitter for more predictable chunking
        logger.debug("Initializing text splitters...")
        # Using SemanticChunker with the Gemini embedding model (kept for backup).
        self.semantic_splitter = SemanticChunker(
            embeddings=gemini_embeddings, 
            breakpoint_threshold_type="percentile"
        )
        
        # Primary text splitter: RecursiveCharacterTextSplitter (more reliable size control)
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=DEFAULT_CHUNK_SIZE,
            chunk_overlap=DEFAULT_CHUNK_OVERLAP,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""]
        )
        
        logger.debug(f"Text splitters initialized.")
        
        self.encoding = tiktoken.get_encoding("cl100k_base") # Retained for token counting if needed elsewhere
        logger.debug("tiktoken encoding cl100k_base loaded.")
        logger.info("DocumentProcessor initialized successfully.")

    async def process_document(self, document_id: int, file_path: str) -> Dict[str, Any]:
        """
        Processes a document: splits into chunks, generates embeddings with contextual headers,
        and saves them to the database including metadata for references.
        """
        logger.info(f"Starting process_document for ID: {document_id}, Path: {file_path}")
        try:
            logger.debug(f"Fetching document name for ID: {document_id}")
            doc_result = self.supabase.table("documents").select("name").eq("id", document_id).execute()
            if doc_result.data and doc_result.data[0].get("name"):
                document_name = doc_result.data[0]["name"]
                logger.debug(f"Document name '{document_name}' found for ID: {document_id}.")
            else:
                document_name = f"Document {document_id}"
                logger.warning(f"Document name not found for ID: {document_id}, using default: '{document_name}'. Response: {doc_result}")
            
            # Use document_name as the initial document title / header
            current_overall_header = document_name 
            
            logger.debug(f"Updating document status to 'processing' for ID: {document_id}")
            await self._update_document_status(document_id, "processing")
            
            # Load and split document using the configured text_splitter
            logger.debug(f"Calling _load_and_split_document for file: {file_path}")
            raw_chunks = await self._load_and_split_document(file_path) # These are Langchain Document objects
            logger.info(f"Document ID: {document_id} ('{document_name}') split into {len(raw_chunks)} raw chunks using {type(self.text_splitter).__name__}.")
            
            # Check if number of chunks exceeds limit
            if len(raw_chunks) > MAX_VECTORS_PER_DOCUMENT:
                logger.warning(f"Number of chunks ({len(raw_chunks)}) exceeds maximum limit ({MAX_VECTORS_PER_DOCUMENT}) for document ID: {document_id}")
                # Trim the chunks to avoid processing too many
                raw_chunks = raw_chunks[:MAX_VECTORS_PER_DOCUMENT]
                logger.info(f"Trimmed chunks to {len(raw_chunks)} for document ID: {document_id}")
            
            # Ensure API key is configured for embedding generation
            api_key = os.environ.get("GEMINI_API_KEY")
            if not api_key:
                logger.error(f"GEMINI_API_KEY not found for document ID: {document_id} processing. This should have been caught earlier.")
                await self._update_document_status(document_id, "failed", note="Missing GEMINI_API_KEY")
                return {"success": False, "error": "Missing GEMINI_API_KEY", "document_id": document_id}
            # genai.configure(api_key=api_key) # Already configured in __init__ and on each embedding call if needed

            processed_chunks_for_db = []
            chunk_meta_info_for_bc = [] # For backward compatibility "embeddings" table

            for i, chunk_doc in enumerate(raw_chunks):
                try:
                    if not chunk_doc.page_content.strip():
                        logger.warning(f"Empty content in chunk {i} from Langchain Splitter (doc ID: {document_id}), skipping.")
                        continue
                        
                    # Determine header and page number from metadata if available
                    header_to_use = chunk_doc.metadata.get('header', f"Chunk {i+1}")
                    page_number = chunk_doc.metadata.get('page_number') # Might be None

                    # Original text (without header) for token count and potentially for embedding if desired
                    original_chunk_text = chunk_doc.page_content
                    
                    # Text to be embedded and stored (potentially with header prepended)
                    # For now, let's assume we embed and store text WITH header, as before.
                    chunk_text_with_header = f"{header_to_use}\n\n{original_chunk_text}" if header_to_use else original_chunk_text

                    # Check if we've reached the vector limit during processing
                    if len(processed_chunks_for_db) >= MAX_VECTORS_PER_DOCUMENT:
                        logger.warning(f"Maximum number of vectors ({MAX_VECTORS_PER_DOCUMENT}) reached during processing for document ID: {document_id}")
                        break

                    embedding_vector = await self._generate_embedding(chunk_text_with_header)
                    if not embedding_vector or len(embedding_vector) == 0:
                        logger.error(f"Empty embedding returned for chunk {i} (doc ID: {document_id}). Original text snippet: {original_chunk_text[:100]}...")
                        continue
                    
                    token_count = len(self.encoding.encode(original_chunk_text)) # Token count of the original text

                    processed_chunks_for_db.append({
                        "document_id": document_id,
                        "chunk_text": chunk_text_with_header,  # Storing text WITH header
                        "embedding": embedding_vector,        # Embedding is of text WITH header
                        "content_token_count": token_count,   # Token count of ORIGINAL text
                        # These were part of the old schema/RPC, ensure SQL function handles them or they are added to RPC params
                        # "chunk_header": header_to_use, 
                        # "page_number": page_number,
                        # "section": header_to_use 
                    })
                    # For backward compatibility with 'embeddings' table
                    chunk_meta_info_for_bc.append({"chunk_index": i, "original_text": original_chunk_text, "header": header_to_use})
                    logger.info(f"Generated embedding for chunk {i+1}/{len(raw_chunks)} (doc ID: {document_id})")
                except Exception as e_emb:
                    logger.error(f"Error generating embedding or preparing chunk {i} (doc ID: {document_id}): {e_emb}", exc_info=True)
                    continue # Skip this chunk
            
            if not processed_chunks_for_db:
                logger.warning(f"No processable chunks with embeddings found for document {document_id} after generation. Update status to failed.")
                await self._update_document_status(document_id, "failed", "No processable chunks with embeddings after generation.")
                return {"success": False, "document_id": document_id, "chunks_created": 0, "error": "No processable chunks with embeddings after generation"}

            # ----- Save chunks to 'document_chunks' table using RPC -----
            logger.info(f"Attempting to save {len(processed_chunks_for_db)} processed chunks to 'document_chunks' via RPC 'insert_document_chunk_basic' for doc ID: {document_id}.")
            successful_rpc_inserts = 0
            any_rpc_insert_failed = False
            first_rpc_error_message = ""

            for i, chunk_data_to_insert in enumerate(processed_chunks_for_db):
                try:
                    logger.debug(f"Preparing to insert chunk {i+1}/{len(processed_chunks_for_db)} via RPC for doc ID: {document_id}. "
                                 f"Snippet: '{chunk_data_to_insert['chunk_text'][:50]}...', Emb Dim: {len(chunk_data_to_insert['embedding']) if chunk_data_to_insert.get('embedding') else 'N/A'}")

                    params_for_rpc = {
                        "p_document_id": chunk_data_to_insert["document_id"],
                        "p_chunk_text": chunk_data_to_insert["chunk_text"],
                        "p_embedding": chunk_data_to_insert["embedding"],
                        "p_content_token_count": chunk_data_to_insert["content_token_count"]
                        # The RPC 'insert_document_chunk_basic' is defined to take these four parameters.
                        # chunk_header, page_number, section are not part of this basic RPC.
                        # If they need to be stored, the SQL function 'insert_document_chunk_basic'
                        # and its parameters in this call would need to be updated.
                    }
                    
                    rpc_response = self.supabase.rpc("insert_document_chunk_basic", params_for_rpc).execute()

                    if hasattr(rpc_response, 'data') and rpc_response.data and len(rpc_response.data) > 0:
                        inserted_chunk_id = rpc_response.data[0].get('id', 'N/A') # RPC returns the inserted row
                        logger.info(f"Successfully inserted chunk {i+1}/{len(processed_chunks_for_db)} via RPC for doc ID: {document_id}. Inserted 'document_chunks' ID: {inserted_chunk_id}")
                        successful_rpc_inserts += 1
                    elif hasattr(rpc_response, 'error') and rpc_response.error:
                        error_msg = f"RPC call error for 'document_chunks' chunk {i+1} (doc ID: {document_id}): Code: {rpc_response.error.code}, Message: {rpc_response.error.message}, Details: {rpc_response.error.details}, Hint: {rpc_response.error.hint}"
                        logger.error(error_msg)
                        any_rpc_insert_failed = True
                        if not first_rpc_error_message: first_rpc_error_message = error_msg
                    else:
                        warn_msg = f"RPC call for 'document_chunks' chunk {i+1} (doc ID: {document_id}) returned no data and no explicit error. Status: {rpc_response.status_code if hasattr(rpc_response, 'status_code') else 'N/A'}. Resp: {rpc_response}"
                        logger.warning(warn_msg)
                        any_rpc_insert_failed = True # Treat as failure for safety
                        if not first_rpc_error_message: first_rpc_error_message = warn_msg
                except Exception as e_rpc_call:
                    error_msg = f"Exception during RPC call for 'document_chunks' chunk {i+1} (doc ID: {document_id}): {e_rpc_call}"
                    logger.error(error_msg, exc_info=True)
                    any_rpc_insert_failed = True
                    if not first_rpc_error_message: first_rpc_error_message = error_msg
            
            logger.info(f"Finished RPC inserts for 'document_chunks' for doc ID: {document_id}. Successfully inserted: {successful_rpc_inserts}/{len(processed_chunks_for_db)}.")

            if any_rpc_insert_failed or (len(processed_chunks_for_db) > 0 and successful_rpc_inserts == 0):
                final_status_reason = first_rpc_error_message or "One or more chunks failed to insert into 'document_chunks' via RPC."
                logger.error(f"Overall 'document_chunks' RPC insertion for doc ID {document_id} failed or partially failed. Error: {final_status_reason}")
                await self._update_document_status(document_id, "failed", final_status_reason)
                return {"success": False, "document_id": document_id, "chunks_created": successful_rpc_inserts, "error": final_status_reason}
            
            # If we reach here, all RPC inserts to 'document_chunks' were successful.
            logger.info(f"All {successful_rpc_inserts} chunks for document ID {document_id} saved successfully to 'document_chunks' via RPC.")
            
            # --- Backward Compatibility: Save to 'embeddings' table ---
            # This uses 'chunk_meta_info_for_bc' and 'processed_chunks_for_db' which should be aligned.
            if successful_rpc_inserts > 0: # Only proceed if primary chunks were saved
                logger.info(f"Attempting to save {len(chunk_meta_info_for_bc)} items to 'embeddings' table for backward compatibility (doc ID: {document_id}).")
                bc_inserts_count = 0
                for i, meta_info in enumerate(chunk_meta_info_for_bc):
                    # Find the corresponding processed chunk data for the embedding
                    # This assumes processed_chunks_for_db[i] corresponds to chunk_meta_info_for_bc[i]
                    if i < len(processed_chunks_for_db):
                        embedding_vector_for_bc = processed_chunks_for_db[i]["embedding"]
                        
                        embedding_data_for_bc_table = {
                            "content": meta_info["original_text"], 
                            "metadata": {"document_id": document_id, "chunk_index": meta_info["chunk_index"], "header": meta_info.get("header")},
                            "embedding": embedding_vector_for_bc,
                        }
                        try:
                            bc_response = self.supabase.table("embeddings").insert(embedding_data_for_bc_table).execute()
                            if hasattr(bc_response, 'data') and bc_response.data:
                                bc_inserts_count += 1
                            elif hasattr(bc_response, 'error') and bc_response.error:
                                logger.error(f"BC Save: Error inserting item {i+1} to 'embeddings' for doc {document_id}: {bc_response.error.message}")
                            # else: log warning for no data/no error if necessary
                        except Exception as e_bc_insert_exc:
                            logger.error(f"BC Save: Exception inserting item {i+1} to 'embeddings' for doc {document_id}: {e_bc_insert_exc}", exc_info=True)
                    else:
                        logger.warning(f"BC Save: Mismatch between chunk_meta_info_for_bc and processed_chunks_for_db at index {i}. Cannot save to 'embeddings' table.")
                
                logger.info(f"Backward compatibility: Saved {bc_inserts_count}/{len(chunk_meta_info_for_bc)} items to 'embeddings' table for doc ID: {document_id}.")
            
            # Update main document status and content (if all primary chunks saved)
            await self._update_document_status(document_id, "completed")
            
            # Only try to update document content if we processed fewer than MAX_VECTORS_PER_DOCUMENT
            if len(raw_chunks) <= MAX_VECTORS_PER_DOCUMENT:
                await self._update_document_content(document_id, raw_chunks)
            else:
                logger.warning(f"Skipping document content update due to large chunk count for document ID: {document_id}")
            
            return {
                "success": True,
                "document_id": document_id,
                "chunks_created": successful_rpc_inserts,
                "total_chunks": len(raw_chunks),
                "trimmed": len(raw_chunks) > MAX_VECTORS_PER_DOCUMENT
            }
            
        except Exception as e:
            logger.error(f"Error in process_document for ID {document_id}: {str(e)}", exc_info=True)
            await self._update_document_status(document_id, "failed", str(e))
            return {"success": False, "document_id": document_id, "chunks_created": 0, "error": str(e)}

    async def _load_and_split_document(self, file_path: str) -> List[Any]:
        """
        Loads a document from file_path and splits it into chunks.
        Returns list of Document objects with metadata.
        """
        logger.debug(f"Starting _load_and_split_document for file: {file_path}")
        
        try:
            # Basic metadata to include with each chunk
            original_metadata = {
                "source": file_path,
                "file_name": Path(file_path).name
            }
            
            # Extract text based on file type
            file_extension = Path(file_path).suffix.lower()
            if file_extension == '.pdf':
                doc_text_content = self._extract_pdf_text(file_path)
            elif file_extension in ['.docx', '.doc']:
                doc_text_content = self._extract_docx_text(file_path)
            elif file_extension in ['.txt', '.md', '.html', '.json']:
                with open(file_path, 'r', encoding='utf-8') as f:
                    doc_text_content = f.read()
                logger.debug(f"Read {len(doc_text_content)} characters from text file: {file_path}")
            else:
                logger.error(f"Unsupported file type: {file_extension} for file {file_path}")
                raise ValueError(f"Unsupported file type: {file_extension}")
            
            if not doc_text_content.strip():
                logger.warning(f"Document {file_path} has no text content after extraction. Returning empty list of chunks.")
                return []

            logger.debug(f"Splitting text content (length: {len(doc_text_content)}) using {type(self.text_splitter).__name__}'s split_documents method...")
            
            # Create a Document object
            doc = type('Document', (), {'page_content': doc_text_content, 'metadata': original_metadata})()
            
            # Use split_documents which returns List[Document]
            chunks = self.text_splitter.split_documents([doc])
            logger.info(f"Successfully split document {file_path} into {len(chunks)} chunks.")

            # Add metadata for each chunk
            for i, chunk in enumerate(chunks):
                chunk.metadata["chunk_index"] = i
                
            return chunks
            
        except Exception as e:
            logger.error(f"Error in _load_and_split_document for {file_path}: {str(e)}", exc_info=True)
            raise

    def _extract_pdf_text(self, file_path: str) -> str:
        """חילוץ טקסט מקובץ PDF"""
        logger.debug(f"Starting _extract_pdf_text for: {file_path}")
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                num_pages = len(pdf_reader.pages)
                logger.debug(f"PDF has {num_pages} pages.")
                for i, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                    logger.debug(f"Extracted text from page {i+1}/{num_pages}. Length: {len(page_text) if page_text else 0}")
            logger.info(f"Finished extracting text from PDF: {file_path}. Total length: {len(text)}")
        except Exception as e:
            logger.error(f"Error extracting PDF text from {file_path}: {str(e)}", exc_info=True)
            raise
        return text

    def _extract_docx_text(self, file_path: str) -> str:
        """חילוץ טקסט מקובץ DOCX"""
        logger.debug(f"Starting _extract_docx_text for: {file_path}")
        try:
            doc = DocxDocument(file_path)
            text = ""
            logger.debug(f"DOCX file has {len(doc.paragraphs)} paragraphs.")
            for i, paragraph in enumerate(doc.paragraphs):
                text += paragraph.text + "\n"
                # Add less verbose logging here, maybe every N paragraphs or by text length
            logger.info(f"Finished extracting text from DOCX: {file_path}. Total length: {len(text)}")
            return text
        except Exception as e:
            logger.error(f"Error extracting DOCX text from {file_path}: {str(e)}", exc_info=True)
            raise

    async def _generate_embedding(self, text: str, is_query: bool = False) -> Optional[List[float]]:
        """יצירת embedding לטקסט"""
        task_type = "retrieval_query" if is_query else "retrieval_document"
        cleaned_text = text[:50].replace('\n', ' ')
        logger.debug(f"Attempting to generate embedding for text (first 50 chars): '{cleaned_text}' with task_type: {task_type}")

        if not text or not text.strip():
            logger.warning(f"Cannot generate embedding for empty or whitespace-only text. Task type: {task_type}")
            return None
        
        # Ensure API key is available (it should be, from __init__ or process_document)
        api_key = os.environ.get("GEMINI_API_KEY")
        if not api_key:
            logger.error("GEMINI_API_KEY not found when trying to generate embedding.")
            return None
        # genai.configure(api_key=api_key) # Potentially redundant if globally configured, but safe.
        
        try:
            # The model name for embeddings might be different from generative models.
            # Using the model specified for embeddings, e.g., "models/embedding-001"
            # Task types: "retrieval_query", "retrieval_document", "semantic_similarity", "classification", "clustering"
            logger.debug(f"Calling genai.embed_content with model models/embedding-001 and task_type {task_type}")
            result = genai.embed_content(
                model="models/embedding-001", # Correct embedding model
                content=text,
                task_type=task_type
            )
            logger.debug(f"Successfully received embedding result from Gemini API for task_type: {task_type}. Embedding length: {len(result['embedding']) if result and 'embedding' in result else 'N/A'}")
            return result["embedding"]
        except Exception as e:
            logger.error(f"Error generating Gemini embedding for task_type {task_type}: {e}", exc_info=True)
            return None

    async def _update_document_status(self, document_id: int, status: str, note: Optional[str] = None):
        """עדכון סטטוס המסמך במסד הנתונים"""
        logger.info(f"Updating status for document ID {document_id} to '{status}'. Note: '{note or ''}'")
        try:
            update_data = {"processing_status": status}
            if note is not None:
                update_data["processing_notes"] = note
            
            response = self.supabase.table("documents").update(update_data).eq("id", document_id).execute()
            # Check response for errors (Supabase client specific)
            if hasattr(response, 'data') and response.data:
                logger.debug(f"Successfully updated status for document ID {document_id} to {status}. Response: {response.data}")
            elif hasattr(response, 'error') and response.error is not None:
                logger.error(f"Supabase error updating status for document ID {document_id}: Code: {response.error.code}, Message: {response.error.message}")
            else:
                logger.warning(f"Unknown Supabase response or no data returned when updating status for document ID {document_id}: {response}")
        except Exception as e:
            logger.error(f"Exception updating document status for ID {document_id} to {status}: {e}", exc_info=True)

    async def _update_document_content(self, document_id: int, chunks: List[Any]):
        """עדכון תוכן המסמך במסד הנתונים עם הטקסט המחובר של ה-chunks"""
        logger.debug(f"Updating document content for ID {document_id} with {len(chunks)} raw chunks.")
        try:
            # Concatenate the page_content of each chunk to form the full document text
            full_text = "\n".join([chunk.page_content for chunk in chunks])
            # Simple length check as a heuristic for content validity
            if len(full_text) < 20: # Arbitrary small number, adjust as needed
                logger.warning(f"Concatenated text for document ID {document_id} is very short (length {len(full_text)}). Check chunking.")

            response = self.supabase.table("documents").update({
                "content": full_text,
                "token_count": len(self.encoding.encode(full_text)) # Update token count based on concatenated raw text
            }).eq("id", document_id).execute()
            
            if hasattr(response, 'data') and response.data:
                logger.info(f"Successfully updated content for document ID {document_id}.")
            elif hasattr(response, 'error') and response.error is not None:
                logger.error(f"Supabase error updating content for document ID {document_id}: {response.error.message}")
            else:
                logger.warning(f"Unknown Supabase response when updating content for document ID {document_id}: {response}")
        except Exception as e:
            logger.error(f"Exception updating document content for ID {document_id}: {e}", exc_info=True)

    async def delete_document_and_all_chunks(self, document_id: int) -> Dict[str, Any]:
        """
        Deletes a document and all its associated chunks from the database.
        First deletes chunks from 'document_chunks', then the document from 'documents'.
        """
        logger.info(f"Attempting to delete document and all its chunks for document_id: {document_id}")
        
        deleted_chunks_count = 0
        try:
            # Delete chunks from 'document_chunks' table
            logger.debug(f"Deleting chunks for document_id: {document_id} from 'document_chunks' table.")
            delete_chunks_response = self.supabase.table("document_chunks").delete().eq("document_id", document_id).execute()
            
            # Check for error in response - updated to handle different response formats
            chunks_error_detected = False
            if hasattr(delete_chunks_response, 'error') and delete_chunks_response.error is not None:
                chunks_error_detected = True
                chunks_error_message = delete_chunks_response.error
            elif isinstance(delete_chunks_response, dict) and delete_chunks_response.get('error'):
                chunks_error_detected = True
                chunks_error_message = delete_chunks_response.get('error')
            
            if chunks_error_detected:
                logger.error(f"Supabase error deleting chunks for document_id {document_id}: {chunks_error_message}")
                # Even if deleting chunks fails, we might still want to attempt deleting the document record or handle differently.
                # For now, we'll report error and stop.
                return {"success": False, "error": f"Supabase error deleting chunks: {chunks_error_message}", "deleted_chunks_count": 0, "document_deleted": False}
            
            if hasattr(delete_chunks_response, 'data') and delete_chunks_response.data:
                deleted_chunks_count = len(delete_chunks_response.data)
                logger.info(f"Successfully deleted {deleted_chunks_count} chunks for document_id: {document_id}.")
            else:
                logger.info(f"No chunks found or an issue occurred while deleting chunks for document_id: {document_id}. Response: {delete_chunks_response}")
                # This might not be an error if the document had no chunks.

            # Delete the document from 'documents' table
            logger.debug(f"Deleting document record for document_id: {document_id} from 'documents' table.")
            delete_document_response = self.supabase.table("documents").delete().eq("id", document_id).execute()

            document_deleted_successfully = False
            
            # Check for error in response - updated to handle different response formats
            doc_error_detected = False
            if hasattr(delete_document_response, 'error') and delete_document_response.error is not None:
                doc_error_detected = True
                doc_error_message = delete_document_response.error
            elif isinstance(delete_document_response, dict) and delete_document_response.get('error'):
                doc_error_detected = True
                doc_error_message = delete_document_response.get('error')
            
            if doc_error_detected:
                logger.error(f"Supabase error deleting document record for document_id {document_id}: {doc_error_message}")
                # If document deletion fails, we should report it. Chunks might have been deleted.
                return {"success": False, "error": f"Supabase error deleting document record: {doc_error_message}", "deleted_chunks_count": deleted_chunks_count, "document_deleted": False}
            
            if hasattr(delete_document_response, 'data') and delete_document_response.data:
                logger.info(f"Successfully deleted document record for document_id: {document_id}.")
                document_deleted_successfully = True
            else:
                logger.info(f"No document record found or an issue occurred while deleting document_id: {document_id} from 'documents'. Response: {delete_document_response}")
                # This might not be an error if the document was already deleted.
                if not deleted_chunks_count: # if no chunks were deleted either, it implies doc might not have existed
                    document_deleted_successfully = True # Consider it successful if nothing was there to delete

            return {"success": True, "deleted_chunks_count": deleted_chunks_count, "document_deleted": document_deleted_successfully}
            
        except Exception as e:
            logger.error(f"Exception in delete_document_and_all_chunks for document_id {document_id}: {e}", exc_info=True)
            return {"success": False, "error": str(e), "deleted_chunks_count": deleted_chunks_count, "document_deleted": False}

    async def delete_document_embeddings(self, document_id: int) -> Dict[str, Any]:
        """
        Deletes all embeddings (chunks) for a given document ID.
        This is used, for example, before reprocessing a document.
        """
        logger.info(f"Attempting to delete all embeddings/chunks for document_id: {document_id}")
        try:
            # Batch delete document chunks from 'document_chunks' table
            logger.info(f"Starting batch deletion of chunks from 'document_chunks' for document_id: {document_id}")
            batch_size = 500  # Define a reasonable batch size
            while True:
                # Fetch a batch of chunk IDs
                chunk_ids_response = self.supabase.table("document_chunks").select("id").eq("document_id", document_id).limit(batch_size).execute()
                
                if not chunk_ids_response.data:
                    logger.info(f"No more chunks found in 'document_chunks' for document_id: {document_id}")
                    break  # No more chunks to delete
                    
                chunk_ids_to_delete = [chunk['id'] for chunk in chunk_ids_response.data]
                
                if not chunk_ids_to_delete:
                    logger.info(f"No chunk IDs to delete in this batch from 'document_chunks' for document_id: {document_id}")
                    break

                logger.info(f"Deleting batch of {len(chunk_ids_to_delete)} chunks from 'document_chunks' for document_id: {document_id}")
                delete_response = self.supabase.table("document_chunks").delete().in_("id", chunk_ids_to_delete).execute()
                
                # Check for error in response - updated to handle different response formats
                error_detected = False
                if hasattr(delete_response, 'error') and delete_response.error:
                    error_detected = True
                    error_message = delete_response.error
                elif isinstance(delete_response, dict) and delete_response.get('error'):
                    error_detected = True
                    error_message = delete_response.get('error')
                
                if error_detected:
                    logger.error(f"Error deleting a batch of chunks from 'document_chunks' for document_id {document_id}: {error_message}")
                    return {"success": False, "message": f"Failed to delete existing chunks from 'document_chunks': {error_message}"}

                if len(chunk_ids_to_delete) < batch_size:
                    logger.info(f"Processed the last batch of chunks from 'document_chunks' for document_id: {document_id}")
                    break
            
            logger.info(f"Successfully finished batch deletion of chunks from 'document_chunks' for document_id: {document_id}.")

            # Also delete from the old 'embeddings' table for cleanup/BC
            # This can remain a single delete operation as it's less critical / likely smaller
            try:
                logger.info(f"Attempting to delete old embeddings from 'embeddings' table for document_id: {document_id}.")
                delete_old_embeddings_response = self.supabase.table("embeddings").delete().eq("document_id", document_id).execute()
                
                # Check for error in response - updated to handle different response formats
                error_detected = False
                if hasattr(delete_old_embeddings_response, 'error') and delete_old_embeddings_response.error:
                    error_detected = True
                    error_message = delete_old_embeddings_response.error
                elif isinstance(delete_old_embeddings_response, dict) and delete_old_embeddings_response.get('error'):
                    error_detected = True
                    error_message = delete_old_embeddings_response.get('error')
                
                if error_detected:
                    logger.warning(f"Error deleting old embeddings from 'embeddings' table for document_id {document_id}: {error_message}")
                else:
                    logger.info(f"Successfully deleted old embeddings from 'embeddings' table for document_id: {document_id}.")
            except Exception as e_old_emb_del:
                logger.warning(f"Exception during deletion from old 'embeddings' table for document_id {document_id}: {e_old_emb_del}")
            
            return {"success": True, "message": "All embeddings/chunks deleted successfully."}
        except Exception as e:
            logger.error(f"General error in delete_document_embeddings for document_id {document_id}: {str(e)}", exc_info=True)
            return {"success": False, "message": str(e)}

    async def search_documents(self, query: str, limit: int = 10, threshold: float = 0.78) -> List[Dict[str, Any]]:
        """חיפוש סמנטי במסמכים"""
        logger.info(f"Starting semantic search with query (first 50 chars): '{query[:50]}', limit: {limit}, threshold: {threshold}")
        try:
            logger.debug("Semantic search: Generating query embedding...")
            query_embedding = await self._generate_embedding(query, is_query=True)
            if not query_embedding:
                logger.error("Semantic search: Failed to generate query embedding.")
                return []
            logger.debug(f"Semantic search: Query embedding generated. Vector length: {len(query_embedding)} Calling Supabase RPC 'match_documents_semantic' with limit={limit}, threshold={threshold}.")
            
            response = self.supabase.rpc(
                "match_documents_semantic",
                {
                "query_embedding": query_embedding,
                "match_threshold": threshold,
                    "match_count": limit,
                },
            ).execute()
            
            if response.data:
                logger.info(f"Semantic search: Supabase RPC 'match_documents_semantic' executed. Found {len(response.data)} results.")
                # Enhanced logging for each result
                for idx, item in enumerate(response.data):
                    cleaned_text = item.get('chunk_text', '')[:100].replace('\n',' ')
                logger.debug(f"  Result {idx+1}: id={item.get('id')}, doc_id={item.get('document_id')}, header='{item.get('chunk_header')}', similarity={item.get('similarity')}, page_no={item.get('page_number')}, text(100)='{cleaned_text}'") 
                return response.data
            else:
                logger.info("Semantic search: No results from Supabase RPC or error in response.")
                # Log error if present
                if hasattr(response, 'error') and response.error:
                    logger.error(f"Supabase RPC error in semantic search: {response.error.message}")
                return []
            
        except Exception as e:
            logger.error(f"Error during semantic search: {e}", exc_info=True)
            return []

    async def hybrid_search(self, query: str, limit: int = 10, threshold: float = 0.78) -> List[Dict[str, Any]]:
        """חיפוש היברידי במסמכים (שילוב של סמנטי ו-keyword)"""
        logger.info(f"Starting HYBRID search with query (first 50 chars): '{query[:50]}', limit: {limit}, threshold: {threshold}")
        try:
            logger.debug("Hybrid search: Generating query embedding...")
            query_embedding = await self._generate_embedding(query, is_query=True)
            if not query_embedding:
                logger.error("Hybrid search: Failed to generate query embedding.")
                return []
            
            # Define weights for combining results
            semantic_weight = 0.6
            keyword_weight = 0.4
            
            # First try to use the database function if it exists
            try:
                logger.debug(f"Hybrid search: Query embedding generated. Attempting to call Supabase RPC 'hybrid_search_documents'...")
                response = self.supabase.rpc(
                    "hybrid_search_documents", 
                    {
                "query_embedding": query_embedding,
                "query_text": query,
                "match_threshold": threshold,
                        "match_count": limit,
                        "semantic_weight": semantic_weight,
                        "keyword_weight": keyword_weight
                    }
                ).execute()
                
                if hasattr(response, 'data') and response.data:
                    logger.info(f"Hybrid search: Supabase RPC 'hybrid_search_documents' executed. Found {len(response.data)} results.")
                    for i, result in enumerate(response.data[:5]):
                        logger.debug(f"  Result {i+1}: id={result.get('id')}, doc_id={result.get('document_id')}, " + 
                                    f"similarity={result.get('similarity')}, combined={result.get('combined_score')}, " +
                                    f"text(100)='{result.get('chunk_text', '')[:100]}'")
                    return response.data
                else:
                    logger.warning("Hybrid search: Supabase RPC 'hybrid_search_documents' returned no data.")
                    return []
                    
            except Exception as e:
                logger.warning(f"Error calling hybrid_search_documents RPC: {e}. Implementing hybrid search locally.")
                
            # Fallback: Implement hybrid search in Python if the database function is not available
            logger.info("Performing hybrid search with local Python implementation...")
            
            # 1. Get semantic search results (using a lower threshold to get more candidates)
            semantic_threshold = threshold * 0.8  # Lower threshold to get more candidates
            semantic_results = await self.search_documents(query, limit=limit*2, threshold=semantic_threshold)
            logger.info(f"Hybrid search (local): Semantic search returned {len(semantic_results)} results.")
            
            # Create semantic_dict for fast lookups
            semantic_dict = {}
            for result in semantic_results:
                result_id = result.get('id')
                if result_id:
                    semantic_dict[result_id] = {
                        'id': result_id,
                        'document_id': result.get('document_id'),
                        'document_name': result.get('document_name', ''),
                        'chunk_text': result.get('chunk_text', ''),
                        'chunk_header': result.get('chunk_header', ''),
                        'page_number': result.get('page_number'),
                        'section': result.get('section', ''),
                        'similarity': result.get('similarity', 0),
                        'text_match_rank': 0,
                        'combined_score': result.get('similarity', 0) * semantic_weight,
                    }
            
            # 2. Get keyword search results
            # Use full-text search with plainto_tsquery (convert query to text search query)
            keyword_search_query = f"""
                SELECT 
                    dc.id,
                    dc.document_id,
                    d.name as document_name,
                    dc.chunk_text,
                    dc.chunk_header,
                    dc.page_number,
                    dc.section,
                    ts_rank_cd(to_tsvector('hebrew', dc.chunk_text), plainto_tsquery('hebrew', '{query}')) as text_match_rank
                FROM document_chunks dc
                JOIN documents d ON d.id = dc.document_id
                WHERE to_tsvector('hebrew', dc.chunk_text) @@ plainto_tsquery('hebrew', '{query}')
                ORDER BY text_match_rank DESC
                LIMIT {limit*2}
            """
            
            try:
                keyword_response = self.supabase.rpc('execute_sql', { 'query': keyword_search_query }).execute()
                keyword_results = []
                
                if hasattr(keyword_response, 'data') and keyword_response.data:
                    keyword_results = keyword_response.data
                else:
                    # Fallback to direct SQL if RPC doesn't work
                    keyword_response = self.supabase.sql(keyword_search_query).execute()
                    if hasattr(keyword_response, 'data') and keyword_response.data:
                        keyword_results = keyword_response.data
            except Exception as e:
                logger.warning(f"Hybrid search (local): Error executing keyword search: {e}")
                keyword_results = []
                
                # Alternative: Get all document chunks and filter locally
                all_chunks = await self._get_all_document_chunks(limit=100)
                
                # Simple token matching
                query_tokens = set(query.lower().split())
                for chunk in all_chunks:
                    if chunk.get('chunk_text'):
                        chunk_text = chunk.get('chunk_text', '').lower()
                        # Compute a simple match score based on token overlap
                        chunk_tokens = set(chunk_text.split())
                        overlap = len(query_tokens.intersection(chunk_tokens))
                        if overlap > 0:
                            match_rank = overlap / len(query_tokens)
                            chunk['text_match_rank'] = match_rank
                            keyword_results.append(chunk)
            
            logger.info(f"Hybrid search (local): Keyword search returned {len(keyword_results)} results.")
            
            # 3. Process keyword results
            combined_results = {}
            # First, copy all semantic results to combined_results
            combined_results.update(semantic_dict)
            
            # Then, process keyword results
            for result in keyword_results:
                result_id = result.get('id')
                if not result_id:
                    continue
                    
                text_match_rank = result.get('text_match_rank', 0)
                keyword_score = text_match_rank * keyword_weight
                
                if result_id in combined_results:
                    # If we already have this result from semantic search, update it
                    combined_results[result_id]['text_match_rank'] = text_match_rank
                    combined_results[result_id]['combined_score'] += keyword_score
                else:
                    # Otherwise add it as a new result
                    combined_results[result_id] = {
                        'id': result_id,
                        'document_id': result.get('document_id'),
                        'document_name': result.get('document_name', ''),
                        'chunk_text': result.get('chunk_text', ''),
                        'chunk_header': result.get('chunk_header', ''),
                        'page_number': result.get('page_number'),
                        'section': result.get('section', ''),
                        'similarity': 0,  # No semantic similarity
                        'text_match_rank': text_match_rank,
                        'combined_score': keyword_score,
                    }
            
            # 4. Finalize results - sort by combined_score and take top 'limit' results
            final_results = list(combined_results.values())
            final_results.sort(key=lambda x: x.get('combined_score', 0), reverse=True)
            final_results = [r for r in final_results if r.get('combined_score', 0) > (threshold * semantic_weight)]
            final_results = final_results[:limit]
            
            logger.info(f"Hybrid search (local): Combined and filtered results. Returning {len(final_results)} results.")
            
            # 5. Log a few results for debugging
            for i, result in enumerate(final_results[:5]):
                logger.debug(f"  Result {i+1}: id={result.get('id')}, doc_id={result.get('document_id')}, " + 
                            f"similarity={result.get('similarity')}, combined={result.get('combined_score')}, " +
                            f"text(100)='{result.get('chunk_text', '')[:100]}'")
            
            return final_results
                
        except Exception as e:
            logger.error(f"Error in hybrid_search: {e}", exc_info=True)
            return []

    async def _get_all_document_chunks(self, limit: int = 100) -> List[Dict[str, Any]]:
        """Get a set of document chunks for local processing"""
        try:
            response = self.supabase.table("document_chunks").select(
                "id,document_id,chunk_text,chunk_header,page_number,section,documents(name)"
            ).limit(limit).execute()
            
            if hasattr(response, 'data') and response.data:
                # Transform the results to match our expected format
                results = []
                for item in response.data:
                    document_name = None
                    if 'documents' in item and item['documents']:
                        document_name = item['documents'].get('name')
                    
                    results.append({
                        'id': item.get('id'),
                        'document_id': item.get('document_id'),
                        'document_name': document_name,
                        'chunk_text': item.get('chunk_text', ''),
                        'chunk_header': item.get('chunk_header', ''),
                        'page_number': item.get('page_number'),
                        'section': item.get('section', '')
                    })
                return results
            return []
        except Exception as e:
            logger.error(f"Error in _get_all_document_chunks: {e}", exc_info=True)
            return [] 